import io
import logging
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    ListItem,
    ListFlowable,
)
from reportlab.lib.enums import TA_LEFT
from bs4 import BeautifulSoup
import re

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """
    Handles generation of different document formats from text content
    """

    @staticmethod
    def generate_docx(text):
        """
        Generate a Word document from text

        Args:
            text (str): The text content to include in the document

        Returns:
            io.BytesIO: A buffer containing the Word document
        """
        try:
            # Create Word document
            doc = Document()

            # Check if content is HTML
            if text.strip().startswith("<"):
                # Parse HTML with BeautifulSoup
                soup = BeautifulSoup(text, "html.parser")

                # Process HTML content
                for element in soup.find_all(
                    ["p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol", "div"]
                ):
                    if element.name.startswith("h"):
                        # Handle headings
                        level = int(element.name[1])
                        p = doc.add_paragraph()
                        p.style = f"Heading {level}"

                        # Process inline formatting within headings
                        for content in element.contents:
                            if hasattr(content, "name"):
                                if content.name in ["strong", "b"]:
                                    p.add_run(content.get_text()).bold = True
                                elif content.name in ["em", "i"]:
                                    p.add_run(content.get_text()).italic = True
                                elif content.name == "u":
                                    p.add_run(content.get_text()).underline = True
                                elif content.name == "span" and content.get("style"):
                                    run = p.add_run(content.get_text())
                                    # Handle color if specified
                                    style = content.get("style", "")
                                    color_match = re.search(
                                        r"color:\s*(#[0-9a-fA-F]{6})", style
                                    )
                                    if color_match:
                                        hex_color = color_match.group(1).lstrip("#")
                                        rgb = tuple(
                                            int(hex_color[i : i + 2], 16)
                                            for i in (0, 2, 4)
                                        )
                                        run.font.color.rgb = RGBColor(*rgb)
                            else:
                                p.add_run(str(content))
                    elif element.name == "p":
                        # Handle paragraphs
                        p = doc.add_paragraph()

                        # Process inline formatting
                        for content in element.contents:
                            if hasattr(content, "name"):
                                if content.name in ["strong", "b"]:
                                    p.add_run(content.get_text()).bold = True
                                elif content.name in ["em", "i"]:
                                    p.add_run(content.get_text()).italic = True
                                elif content.name == "u":
                                    p.add_run(content.get_text()).underline = True
                                elif content.name == "br":
                                    p.add_run("\n")
                                elif content.name == "span" and content.get("style"):
                                    run = p.add_run(content.get_text())
                                    # Handle color if specified
                                    style = content.get("style", "")
                                    color_match = re.search(
                                        r"color:\s*(#[0-9a-fA-F]{6})", style
                                    )
                                    if color_match:
                                        hex_color = color_match.group(1).lstrip("#")
                                        rgb = tuple(
                                            int(hex_color[i : i + 2], 16)
                                            for i in (0, 2, 4)
                                        )
                                        run.font.color.rgb = RGBColor(*rgb)
                            else:
                                p.add_run(str(content))
                    elif element.name == "ul" or element.name == "ol":
                        # Handle lists
                        for li in element.find_all("li", recursive=False):
                            p = doc.add_paragraph(
                                style=(
                                    "List Bullet"
                                    if element.name == "ul"
                                    else "List Number"
                                )
                            )

                            # Process inline formatting within list items
                            for content in li.contents:
                                if hasattr(content, "name"):
                                    if content.name in ["strong", "b"]:
                                        p.add_run(content.get_text()).bold = True
                                    elif content.name in ["em", "i"]:
                                        p.add_run(content.get_text()).italic = True
                                    elif content.name == "u":
                                        p.add_run(content.get_text()).underline = True
                                    elif content.name == "span" and content.get(
                                        "style"
                                    ):
                                        run = p.add_run(content.get_text())
                                        # Handle color if specified
                                        style = content.get("style", "")
                                        color_match = re.search(
                                            r"color:\s*(#[0-9a-fA-F]{6})", style
                                        )
                                        if color_match:
                                            hex_color = color_match.group(1).lstrip("#")
                                            rgb = tuple(
                                                int(hex_color[i : i + 2], 16)
                                                for i in (0, 2, 4)
                                            )
                                            run.font.color.rgb = RGBColor(*rgb)
                                else:
                                    p.add_run(str(content))
            else:
                # Process plain text
                for paragraph in text.split("\n"):
                    doc.add_paragraph(paragraph)

            # Save to memory buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return buffer
        except Exception as e:
            logger.error(f"Error generating Word document: {e}")
            raise

    @staticmethod
    def generate_pdf(text):
        """
        Generate a PDF document from text

        Args:
            text (str): The text content to include in the document

        Returns:
            io.BytesIO: A buffer containing the PDF document
        """
        try:
            buffer = io.BytesIO()

            # Create a document template with proper margins
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=72,  # 1 inch
                leftMargin=72,  # 1 inch
                topMargin=72,  # 1 inch
                bottomMargin=72,  # 1 inch
            )

            # Create styles
            styles = getSampleStyleSheet()
            normal_style = ParagraphStyle(
                "NormalWithSpacing",
                parent=styles["Normal"],
                fontName="Helvetica",
                fontSize=11,
                leading=14,  # Line height
                alignment=TA_LEFT,
                spaceAfter=10,
            )

            # Process the text into paragraphs
            story = []

            # Check if content is HTML
            if text.strip().startswith("<"):
                # Parse HTML with BeautifulSoup
                soup = BeautifulSoup(text, "html.parser")

                # Process HTML content
                for element in soup.find_all(
                    ["p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol", "div"]
                ):
                    if element.name.startswith("h"):
                        # Handle headings
                        level = int(element.name[1])
                        heading_style = (
                            styles[f"Heading{level}"]
                            if f"Heading{level}" in styles
                            else styles["Heading1"]
                        )

                        # Clean the HTML for ReportLab
                        clean_html = DocumentGenerator._clean_html_for_reportlab(
                            str(element)
                        )

                        p = Paragraph(clean_html, heading_style)
                        story.append(p)
                    elif element.name == "p":
                        # Handle paragraphs - clean HTML for ReportLab
                        clean_html = DocumentGenerator._clean_html_for_reportlab(
                            str(element)
                        )

                        p = Paragraph(clean_html, normal_style)
                        story.append(p)
                    elif element.name == "ul" or element.name == "ol":
                        # Handle lists
                        list_items = []
                        for li in element.find_all("li", recursive=False):
                            # Clean HTML for ReportLab
                            clean_html = DocumentGenerator._clean_html_for_reportlab(
                                str(li)
                            )

                            list_items.append(
                                ListItem(Paragraph(clean_html, normal_style))
                            )

                        list_style = (
                            "bulletlist" if element.name == "ul" else "orderedlist"
                        )
                        story.append(ListFlowable(list_items, style=list_style))

                    # Add space after each element
                    story.append(Spacer(1, 6))
            else:
                # Process plain text
                for line in text.split("\n"):
                    # Skip empty lines but add space
                    if not line.strip():
                        story.append(Spacer(1, 12))
                        continue

                    # Add the paragraph with proper encoding
                    p = Paragraph(
                        line.replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;"),
                        normal_style,
                    )
                    story.append(p)

            # Build the document
            doc.build(story)

            buffer.seek(0)
            return buffer
        except Exception as e:
            logger.error(f"Error generating PDF document: {e}")
            raise

    @staticmethod
    def _clean_html_for_reportlab(html_content):
        """
        Clean HTML content to make it compatible with ReportLab's XML parser

        Args:
            html_content (str): The HTML content to clean

        Returns:
            str: Cleaned HTML content
        """
        # Parse the HTML with BeautifulSoup
        soup = BeautifulSoup(html_content, "html.parser")

        # Process all tags with style attributes
        for tag in soup.find_all(lambda tag: tag.has_attr("style")):
            # Extract the style attribute
            style = tag.get("style", "")

            # Remove the style attribute as ReportLab doesn't support it
            del tag["style"]

            # Extract color information if present
            color_match = re.search(r"color:\s*rgb\((\d+),\s*(\d+),\s*(\d+)\)", style)
            if color_match:
                r, g, b = map(int, color_match.groups())
                # Convert RGB to hex
                hex_color = f"#{r:02x}{g:02x}{b:02x}"
                # Add color attribute that ReportLab understands
                tag["color"] = hex_color

        # Convert specific tags to ReportLab format
        html_content = str(soup)

        # Replace HTML tags with ReportLab-compatible tags
        html_content = html_content.replace("<strong>", "<b>").replace(
            "</strong>", "</b>"
        )
        html_content = html_content.replace("<em>", "<i>").replace("</em>", "</i>")

        # Remove any remaining style attributes (in case BeautifulSoup didn't catch them)
        html_content = re.sub(r' style="[^"]*"', "", html_content)

        return html_content
